"""
生成Word文档 - 每页两个发票
"""
import os
import sys
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from io import BytesIO

# PDF转图片支持
try:
    from pdf2image import convert_from_path
    PDF_SUPPORT = True
except:
    PDF_SUPPORT = False


def add_invoice_to_doc(doc, image_path, info_text):
    """添加发票到文档"""
    try:
        # 处理PDF
        if Path(image_path).suffix.lower() == ".pdf":
            if not PDF_SUPPORT:
                print(f"   [SKIP] PDF支持未安装")
                return False

            # 转换PDF第一页为图片
            project_root = os.path.dirname(os.path.abspath(__file__))
            poppler_path = os.path.join(project_root, "poppler-24.08.0", "Library", "bin")
            images = convert_from_path(image_path, first_page=1, last_page=1, dpi=150, poppler_path=poppler_path)

            if not images:
                print(f"   [SKIP] PDF转换失败")
                return False

            # 将PIL图片保存到临时缓冲区
            img_buffer = BytesIO()
            images[0].save(img_buffer, format='PNG')
            img_buffer.seek(0)

            # 添加图片到文档
            doc.add_picture(img_buffer, width=Inches(5.5))
        else:
            # 直接添加图片
            doc.add_picture(image_path, width=Inches(5.5))

        # 添加发票信息
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(info_text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)

        return True

    except Exception as e:
        print(f"   [ERROR] 添加失败: {e}")
        return False


def generate_word_document(output_folder, output_filename="发票汇总.docx"):
    """生成Word文档"""
    print("="*60)
    print("生成Word文档")
    print("="*60)

    # 检查输出文件夹
    if not os.path.exists(output_folder):
        print(f"[ERROR] 文件夹不存在: {output_folder}")
        return

    # 读取统计信息
    stats_file = os.path.join(output_folder, "statistics.json")
    if not os.path.exists(stats_file):
        print(f"[ERROR] 未找到统计文件: {stats_file}")
        return

    with open(stats_file, "r", encoding="utf-8") as f:
        stats = json.load(f)

    # 获取所有发票文件
    supported_formats = [".jpg", ".jpeg", ".png", ".pdf"]
    invoice_files = []
    for file in os.listdir(output_folder):
        if Path(file).suffix.lower() in supported_formats:
            invoice_files.append(os.path.join(output_folder, file))

    if not invoice_files:
        print("[ERROR] 未找到发票文件")
        return

    # 按文件名排序
    invoice_files.sort()

    print(f"\n找到 {len(invoice_files)} 个发票")

    # 创建Word文档
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4高度
    section.page_width = Inches(8.27)    # A4宽度
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # 添加标题
    title = doc.add_heading('发票汇总报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加日期
    from datetime import datetime
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(f"生成日期: {datetime.now().strftime('%Y年%m月%d日')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # 空行

    # 添加统计信息表格
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'

    # 表格内容
    table.cell(0, 0).text = "发票总数"
    table.cell(0, 1).text = f"{stats['total']} 个"
    table.cell(1, 0).text = "总金额"
    table.cell(1, 1).text = f"{stats['total_amount']:.2f} 元"

    # 分类统计
    type_texts = []
    for inv_type, data in stats['by_type'].items():
        type_texts.append(f"{inv_type}: {data['count']}个")
    table.cell(2, 0).text = "分类统计"
    table.cell(2, 1).text = " | ".join(type_texts)

    # 金额分布
    amount_texts = []
    for inv_type, data in stats['by_type'].items():
        amount_texts.append(f"{inv_type}: {data['amount']:.2f}元")
    table.cell(3, 0).text = "金额分布"
    table.cell(3, 1).text = " | ".join(amount_texts)

    doc.add_paragraph()  # 空行
    doc.add_paragraph()  # 空行

    # 添加发票明细标题
    detail_title = doc.add_heading('发票明细', 2)
    detail_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # 空行

    # 每页添加两个发票
    for idx, file_path in enumerate(invoice_files, 1):
        filename = os.path.basename(file_path)
        print(f"\n[{idx}/{len(invoice_files)}] {filename}")

        # 解析文件名获取信息
        name_parts = Path(filename).stem.split('_')
        if len(name_parts) >= 2:
            amount = name_parts[0]
            inv_type = name_parts[1]
            date = name_parts[2] if len(name_parts) >= 3 else ""

            if date and len(date) == 8:
                # 格式化日期
                date_str = f"{date[:4]}-{date[4:6]}-{date[6:]}"
            else:
                date_str = date

            info_text = f"【{idx}】金额: {amount}元 | 类型: {inv_type}"
            if date_str:
                info_text += f" | 日期: {date_str}"
        else:
            info_text = f"【{idx}】{filename}"

        # 添加发票到文档
        success = add_invoice_to_doc(doc, file_path, info_text)

        if success:
            print(f"   [OK] 已添加")

        # 每两个发票后添加分页符（除了最后一个）
        if idx % 2 == 0 and idx < len(invoice_files):
            doc.add_page_break()
        elif idx % 2 == 1 and idx < len(invoice_files):
            # 添加分隔线
            doc.add_paragraph("─" * 60)

    # 保存文档
    output_path = os.path.join(output_folder, output_filename)
    doc.save(output_path)

    print("\n" + "="*60)
    print(f"Word文档已生成: {output_path}")
    print("="*60)


def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("用法: python generate_word.py <输出文件夹路径> [Word文件名]")
        return

    output_folder = sys.argv[1]
    output_filename = sys.argv[2] if len(sys.argv) > 2 else "发票汇总.docx"

    generate_word_document(output_folder, output_filename)


if __name__ == "__main__":
    main()
